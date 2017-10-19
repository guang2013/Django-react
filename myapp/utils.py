#!/usr/bin/env
# -*- coding: utf-8 -*-
import codecs
import os
import django
import docx
os.environ.setdefault('DJANGO_SETTINGS_MODULE','geography.settings')
django.setup()
from myapp.models import MyData


def import_data():
    doc = docx.Document('./input.txt.docx')
    paras = doc.paragraphs
    data = []
    for index,p in enumerate(paras):
        if index!=0 :
            item = p.text.split('\t')
            data.append(MyData(
                sid=item[0],
                description=item[1],
                datetime=item[2],
                longitude=item[3],
                latitude=item[4],
                elevation=item[5],
            ) )
    MyData.objects.bulk_create(data)
    print('success to import data')
    # print('\n'.join(data))

def export_data():
    document = docx.Document()


    items = MyData.objects.all()
    for item in items:
        pass # 


if __name__ == '__main__':
    import_data()
    # export_data()